from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

OUT = Path(r"C:\Users\zhini\Dropbox\ag_banking_one_big_beauty\docs\ag_banking_obbba_research_summary.docx")

doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(0.8)
sec.bottom_margin = Inches(0.8)
sec.left_margin = Inches(1.0)
sec.right_margin = Inches(1.0)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for name, size, color, before, after in [
    ("Heading 1", 16, "2E74B5", 18, 10),
    ("Heading 2", 13, "2E74B5", 14, 7),
    ("Heading 3", 12, "1F4D78", 10, 5),
]:
    st = styles[name]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

bullet_style = styles["List Bullet"]
bullet_style.font.name = "Calibri"
bullet_style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
bullet_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
bullet_style.font.size = Pt(11)
bullet_style.paragraph_format.left_indent = Inches(0.375)
bullet_style.paragraph_format.first_line_indent = Inches(-0.188)
bullet_style.paragraph_format.space_after = Pt(4)
bullet_style.paragraph_format.line_spacing = 1.25

def shade_paragraph(paragraph, fill="F4F6F9"):
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after = Pt(4)
r = title.add_run("Agricultural Banking and the One Big Beautiful Bill Act")
r.font.name = "Calibri"
r._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
r._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
r.font.size = Pt(22)
r.font.bold = True
r.font.color.rgb = RGBColor.from_string("0B2545")

subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(12)
sr = subtitle.add_run("Data, research questions, methods, and current results")
sr.font.size = Pt(11)
sr.font.italic = True
sr.font.color.rgb = RGBColor.from_string("555555")

note = doc.add_paragraph()
note.paragraph_format.left_indent = Inches(0.12)
note.paragraph_format.right_indent = Inches(0.12)
note.paragraph_format.space_before = Pt(2)
note.paragraph_format.space_after = Pt(10)
note.add_run("Scope note. ").bold = True
note.add_run(
    "The deposit and historical lending regressions are empirical estimates. "
    "The Bellman counterfactuals are coarse-grid structural calibrations, not a completed Wang-style simulated-minimum-distance estimation."
)
shade_paragraph(note)

doc.add_heading("1. Research Questions", level=1)
for text in [
    "How do OBBBA increases in ARC/PLC support and eligible base acres change government payments received in agricultural counties?",
    "How much of an increase in total FSA payments remains in local bank deposits, rather than being immediately consumed or invested?",
    "How do the policy payments affect agricultural loan demand, agricultural loan rates, default risk, bank capital, and wholesale funding?",
    "Do banks convert additional deposits into agricultural loans, or do they replace wholesale funding, purchase securities, or retain liquidity?",
    "How much of the policy benefit is retained by banks through deposit and agricultural-loan market power?",
    "How do results change under competitive deposit pricing, competitive agricultural-loan pricing, both markets competitive, and no bank capital constraint?",
    "Do higher Marketing Assistance Loan rates crowd out commercial-bank borrowing while improving farmers' access to government-supported credit?",
    "What are the implications for farmers: lower borrowing needs, lower credit risk, cheaper credit, or redistribution of benefits toward banks?",
]:
    add_bullet(text)

doc.add_heading("2. Datasets Downloaded and Constructed", level=1)
datasets = [
    ("FSA annual payment files, 2014-2025. ", "104 official recipient/payment workbooks (approximately 3.5 GB). Recipient records were aggregated to county, year, and program; names and addresses remain outside Git. Programs were classified as ARC/PLC, MFP, CFAP, other FSA, and total FSA payments."),
    ("FSA ARC/PLC program data. ", "Official 2025 county-crop enrolled base acres, 2025 county PLC yields, projected 2026 effective prices and PLC payment rates, 2026 market-year-average price information, and 2026 ARC-CO county benchmark revenues."),
    ("BEA county farm-income data. ", "CAINC45 county government payments and farm/personal income. The banking overlap is 1994-2022 because BEA discontinued the detailed CAINC45 table after 2022; it was not artificially extended."),
    ("FDIC Summary of Deposits, 1994-2025. ", "Branch-level deposits were aggregated to county markets and used to construct bank service-area weights from prior-year branch deposit shares."),
    ("FDIC Call Reports, 1994-2025. ", "Bank deposits, agricultural production loans, total agricultural loans, total loans, equity, capital ratios, net charge-offs, interest income and expense, reserves, wholesale/nondeposit funding, repos, dividends, taxes, employment, premises expense, and other balance-sheet variables."),
    ("USDA ERS agricultural finance data. ", "Annual agricultural debt by lender, including commercial banks, the Farm Credit System, the Farm Service Agency, and other lenders; these series define the outside-credit environment."),
    ("USDA ERS farm income and wealth data. ", "National net cash farm income was converted to real terms and used to construct farm-income innovations and strong, normal, and downturn states."),
    ("Federal Reserve/FRED data. ", "Daily effective federal funds rate observations were aggregated to annual funding-rate states for the dynamic bank model."),
    ("USDA ERS Marketing Assistance Loan rates. ", "Pre-OBBBA and OBBBA statutory commodity loan rates were recorded for barley, corn, soybeans, wheat, rice, pulses, cotton, peanuts, and other covered crops. Most rates increase by about 10 percent, with heterogeneous increases across commodities."),
    ("Constructed analytical panels. ", "County payment-retention panel, bank service-area policy-exposure panel, deposit/total-loan/agricultural-production-loan BLP panels, dynamic bank-model panel, OBBBA county payment simulations, bank deposit shocks, default sensitivities, and Marketing Assistance Loan exposure indices."),
]
for prefix, body in datasets:
    add_bullet(prefix + body, bold_prefix=prefix)

doc.add_heading("3. Methods", level=1)
methods = [
    ("Payment timing and geography. ", "FSA disbursements from July through December are assigned to the following June SOD observation; January through June payments are assigned to the same June. County payments are not assigned directly to individual banks. Bank exposure is a service-area measure based on prior-year SOD county deposit weights."),
    ("Payment-retention regression. ", "County deposit growth is regressed on total FSA payments divided by lagged county deposits, with county and year fixed effects and standard errors clustered by county. Separate ARC/PLC, MFP, and CFAP regressions are retained as diagnostics."),
    ("OBBBA payment simulation. ", "Baseline and OBBBA ARC/PLC payments are calculated from county-crop base acres, PLC yields, projected effective prices, reference prices, ARC benchmarks, guarantee rates, payment caps, and a proportional scenario for the additional 30 million base acres. ARC outcomes are evaluated over revenue scenarios from 70 to 100 percent of benchmark revenue."),
    ("Deposit counterfactual. ", "The simulated increase in total FSA payments is multiplied by the statistically significant total-FSA deposit-retention coefficient. County deposit growth is mapped to banks using branch-deposit geography."),
    ("BLP market-power estimation. ", "Separate differentiated-product demand systems are estimated for deposits, total loans, and agricultural production loans. Bank and year fixed effects are included; salary expense/assets and premises expense/assets serve as the primary Wang-style supply-cost instruments. Demand derivatives are combined with ownership matrices to recover markups or markdowns."),
    ("Agricultural-lending reduced form. ", "Bank agricultural-loan balances and rates are related to total-FSA payment intensity in the bank's service area using bank and year fixed effects. This estimates the historical net association but does not separately identify borrower demand, bank supply, default risk, or program targeting."),
    ("Balance-sheet channel regressions. ", "Net charge-offs, wholesale funding, reserves, capital, agricultural-loan markups, deposit markdowns, and deposit rates are related to total-FSA service-area intensity. Effects that are not statistically identified are not treated as causal estimates."),
    ("Default-risk sensitivity. ", "Because historical payments are targeted to weak agricultural conditions, the charge-off regression cannot identify the causal insurance effect. The Bellman exercise therefore reports no default effect, a 10 percent charge-off reduction, and a 25 percent charge-off reduction."),
    ("Government-credit substitution. ", "OBBBA Marketing Assistance Loan rate changes are mapped to county crop exposure using enrolled base acres as a proxy. Because MAL eligibility follows current production rather than base acres, this is an exposure calibration, not a forecast of government loan volume."),
    ("Corrected Bellman model. ", "Banks choose loan rates, deposit rates, dividends, securities, reserves, and wholesale funding subject to deposit and loan demand, defaults, balance-sheet feasibility, retained earnings, reserve requirements, and capital constraints. The invalid earlier competition shortcut that multiplied demand slopes by 100 was removed."),
    ("Structural counterfactuals. ", "The model solves no policy; full OBBBA under current market power; funding-only, borrower-liquidity-only, default-only, and government-credit-only channels; alternative default assumptions; competitive deposits; competitive lending; both markets competitive; and no capital constraint."),
    ("Competition restrictions. ", "Competitive agricultural lending constrains the loan rate to risk-adjusted marginal cost. Competitive deposit pricing uses the avoided marginal wholesale funding cost net of deposit servicing and reserve carry. The current implementation uses a two-point bank-state grid and reports convergence diagnostics."),
]
for prefix, body in methods:
    add_bullet(prefix + body, bold_prefix=prefix)

doc.add_heading("4. Current Results", level=1)
results = [
    ("BLP market power. ", "The median deposit markdown is approximately 16.1 basis points; the median total-loan markup is 100.9 basis points; and the median agricultural-production-loan markup is 156.0 basis points."),
    ("Total-FSA deposit response. ", "The total-FSA coefficient is 0.07616 (standard error 0.02337; p = 0.0011). Interpreted mechanically, one additional dollar of FSA disbursement is associated with about 7.6 cents of additional county deposit stock. The separate ARC/PLC, MFP, and CFAP coefficients are not statistically significant and are not used in the main simulation."),
    ("Central OBBBA payment simulation. ", "At actual county revenue equal to 85 percent of ARC benchmark revenue, incremental OBBBA support is approximately $9.32 billion, a 61.2 percent increase relative to the 2025 total-FSA baseline of $15.23 billion."),
    ("Predicted deposits. ", "Applying the total-FSA coefficient predicts approximately $710 million in additional county deposits. The bank service-area mapping produces approximately $729 million because county growth rates are applied to banks' full deposit bases."),
    ("Historical agricultural lending response. ", "Total-FSA service-area intensity is associated with lower agricultural-production loan balances (coefficient -0.7618; p = 0.0098). The agricultural-loan-rate coefficient is positive but statistically insignificant (p = 0.232), so the project does not conclude that OBBBA raises loan rates."),
    ("Reduced-form OBBBA lending prediction. ", "The historical lending relationship predicts agricultural lending declines of approximately 1.21 percent at agricultural banks and 0.23 percent at nonagricultural banks, about $597 million combined. This is consistent with payments reducing farmers' short-term operating-credit needs."),
    ("Default risk. ", "The historical charge-off coefficient is not statistically identified (p = 0.188). The model therefore treats lower default risk as a sensitivity rather than an estimated effect. The 10 percent and 25 percent cases reduce average charge-offs by approximately 8.2 and 20.4 basis points, respectively."),
    ("Funding and capital. ", "Total-FSA exposure is associated with lower wholesale funding (coefficient -0.0335; p = 0.065) and higher bank capital (coefficient 0.0390; p < 0.001). This supports a funding-substitution channel but does not imply that all new deposits become loans."),
    ("Combined Bellman calibration under current market power. ", "For agricultural banks, the combined policy calibration reduces new agricultural loans by 6.64 percent and lowers loan rates by 14.6 basis points. For nonagricultural banks, new agricultural loans decline by 9.36 percent and the loan-rate change is approximately zero."),
    ("Competitive deposit calibration. ", "Competitive deposit pricing raises depositor rates by roughly 138-150 basis points in the coarse model but does not materially change agricultural lending relative to the current-power OBBBA case. The additional funding is largely absorbed by securities and balance-sheet adjustments."),
    ("Competitive lending calibration. ", "Constraining agricultural loan rates to risk-adjusted marginal cost produces much larger lending and lower loan rates: approximately +180 percent and -168 basis points for agricultural banks, and +116 percent and -91 basis points for nonagricultural banks."),
    ("Capital constraint calibration. ", "Removing the capital requirement produces very large lending expansions in the coarse model. This shows that capital can be quantitatively important, but the magnitude is not suitable as a headline estimate."),
    ("Farmer policy implication. ", "OBBBA can improve farmer liquidity and repayment capacity while simultaneously reducing demand for short-term commercial operating loans. Whether farmers also receive cheaper or more available bank credit depends on default-risk pass-through, commercial-bank market power, capital constraints, and competition from FSA/Farm Credit and Marketing Assistance Loans."),
    ("Interpretation limitation. ", "The Bellman accounting identities hold to numerical precision and no solved policy violates the imposed capital constraint. However, competitor-rate fixed-point gaps remain about 0.002-0.012. The structural magnitudes are calibration results and require a finer grid, full equilibrium convergence, and valid instruments for default and outside-credit elasticities before publication."),
]
for prefix, body in results:
    add_bullet(prefix + body, bold_prefix=prefix)

doc.core_properties.title = "Agricultural Banking and OBBBA Research Summary"
doc.core_properties.subject = "Data, methods, research questions, and current results"
doc.core_properties.author = "AG Banking Research Project"

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
