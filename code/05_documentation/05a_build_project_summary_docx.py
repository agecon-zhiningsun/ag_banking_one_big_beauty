from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

OUT = Path(r"C:\Users\zhini\Dropbox\ag_banking_one_big_beauty\docs\05g_ag_banking_obbba_research_summary.docx")

doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(0.8)
sec.bottom_margin = Inches(0.8)
sec.left_margin = Inches(1.0)
sec.right_margin = Inches(1.0)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for name, size, color, before, after in [
    ("Heading 1", 16, "2E74B5", 18, 10),
    ("Heading 2", 13, "2E74B5", 14, 7),
    ("Heading 3", 12, "1F4D78", 10, 5),
]:
    st = styles[name]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

bullet_style = styles["List Bullet"]
bullet_style.font.name = "Calibri"
bullet_style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
bullet_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
bullet_style.font.size = Pt(11)
bullet_style.paragraph_format.left_indent = Inches(0.375)
bullet_style.paragraph_format.first_line_indent = Inches(-0.188)
bullet_style.paragraph_format.space_after = Pt(4)
bullet_style.paragraph_format.line_spacing = 1.25

def shade_paragraph(paragraph, fill="F4F6F9"):
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after = Pt(4)
r = title.add_run("How Government Farm Support Transmits Through Banks: Farmer Welfare, Bank Risk, and a Case Study of the One Big Beautiful Bill Act")
r.font.name = "Calibri"
r._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
r._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
r.font.size = Pt(22)
r.font.bold = True
r.font.color.rgb = RGBColor.from_string("0B2545")

subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(12)
sr = subtitle.add_run("Research framing, literature, data, methods, and current predictions")
sr.font.size = Pt(11)
sr.font.italic = True
sr.font.color.rgb = RGBColor.from_string("555555")

note = doc.add_paragraph()
note.paragraph_format.left_indent = Inches(0.12)
note.paragraph_format.right_indent = Inches(0.12)
note.paragraph_format.space_before = Pt(2)
note.paragraph_format.space_after = Pt(10)
note.add_run("Scope note. ").bold = True
note.add_run(
    "The deposit and historical lending regressions are empirical estimates. "
    "The Bellman counterfactuals are coarse-grid structural calibrations, not a completed Wang-style simulated-minimum-distance estimation."
)
shade_paragraph(note)

doc.add_heading("1. Research Question and Scope", level=1)
lead = doc.add_paragraph()
lead.add_run("Main research question. ").bold = True
lead.add_run(
    "How do government farm-support payments affect farmer welfare and bank risk "
    "when transmitted through imperfectly competitive agricultural banking markets?"
)
shade_paragraph(lead, fill="EAF2F8")

doc.add_paragraph(
    "Farmers and other eligible agricultural producers receive statutory support payments directly; "
    "banks are not program beneficiaries. The banking channel begins only after a payment changes "
    "borrower liquidity, deposit balances, operating-credit demand, repayment risk, or the value of "
    "government-supported credit outside commercial banks. The paper therefore asks how the initial "
    "farmer benefit is propagated, attenuated, or redistributed by the existing banking market."
)

doc.add_paragraph(
    "We use the agricultural provisions of OBBBA as a large, plausibly measurable change in government "
    "support. We combine statutory program changes, county agricultural exposure, historical payment "
    "responses, and estimated banking-market power to evaluate the short- and long-run transmission of "
    "farm support. The OBBBA results are ex ante counterfactual predictions disciplined by historical "
    "relationships, not realized post-policy treatment effects."
)

for text in [
    "Farmer channel: Do support payments reduce costly operating-credit needs and repayment risk, or change future access to and pricing of credit?",
    "Bank channel: Do payment-induced deposits and safer borrowers lead banks to expand agricultural lending, replace wholesale funding, purchase securities, or retain liquidity and capital?",
    "Incidence under current market power: How do estimated deposit markdowns and agricultural-loan markups shape the indirect pass-through of support that farmers initially receive?",
]:
    add_bullet(text)

doc.add_heading("2. Focused Literature Review", level=1)
doc.add_heading("Government support and agricultural markets", level=2)
doc.add_paragraph(
    "A broad agricultural-policy literature studies how support changes production incentives, commodity "
    "marketing, prices, and expectations. Hennessy (1998) shows that income support can affect production "
    "through risk even when payments appear decoupled. Janzen, Swearingen, and Yu (2023) find that Market "
    "Facilitation Program payments allowed grain producers to delay sales and increase storage, shifting "
    "commodity supply across time. Sun's AEPP study complements this work by showing that pre-official "
    "MFP and CFAP announcements changed corn and soybean futures volatility even when official announcements "
    "had little additional effect. These studies establish that support can change expectations and real "
    "decisions before or beyond the mechanical income transfer."
)

doc.add_heading("Distribution, political incidence, and use of payments", level=2)
doc.add_paragraph(
    "A second literature asks who receives support and how recipients use it. Giri, Subedi, and Kassel "
    "(2022) document heterogeneous MFP and CFAP payments across producer groups, while Janzen, Malone, "
    "Schaefer, and Scheitrum (2023) study the political returns to ad hoc farm payments. Research on land "
    "incidence shows that the statutory recipient need not retain the entire economic benefit: Roberts, "
    "Kirwan, and Hopkins (2003), Goodwin, Mishra, and Ortalo-Magne (2003), and Kirwan (2009) examine "
    "capitalization into land values and rents. USDA research further emphasizes that decoupled payments "
    "can raise household welfare, savings, collateral, or investment when farmers face liquidity and "
    "credit constraints. Consequently, a payment can be consumed, saved, invested, used to repay debt, "
    "or transferred to landowners; an increase in local deposits is an empirical question rather than an assumption."
)

doc.add_heading("Agricultural credit, repayment risk, and banks", level=2)
doc.add_paragraph(
    "Agricultural-finance research links household income, collateral, and borrower risk to credit use. "
    "Key (2020) finds that greater household income relaxes credit constraints and increases farm borrowing "
    "and investment. Walraven and Barry (2003) show that commercial banks charge materially higher rates "
    "on agricultural loans they classify as riskier. Recent work directly relates ARC, PLC, MFP, and CFAP "
    "payments to non-real-estate farm debt and delinquency, finding that program effects differ across debt "
    "and delinquency horizons. This literature motivates the borrower-liquidity and default-risk channels, "
    "but it generally does not jointly model how banks reallocate deposits, wholesale funding, securities, "
    "capital, and agricultural credit."
)

doc.add_heading("Financial intermediation and market power", level=2)
doc.add_paragraph(
    "The banking literature shows that deposit and loan market power can alter the pass-through of external "
    "shocks. Wang, Whited, Wu, and Xiao develop a dynamic bank model in which deposit market power, loan "
    "market power, capital regulation, and reserve requirements jointly determine lending responses. This "
    "project applies that intermediation logic to fiscal farm support: the government pays eligible producers, "
    "and banks subsequently choose how to price deposits and loans and allocate their balance sheets under "
    "the market power estimated in agricultural banking data."
)

doc.add_heading("3. Contribution", level=1)
for text in [
    "Adds a financial-intermediation channel to research that has primarily examined commodity markets, political allocation, land incidence, and farmers' direct use of support payments.",
    "Separates the direct beneficiary from the indirect intermediary: farmers receive the payment, while banks respond to changes in deposits, credit demand, repayment risk, and government-credit outside options.",
    "Combines county policy exposure, historical payment-deposit and payment-credit relationships, estimated agricultural-bank market power, and a dynamic bank balance-sheet model.",
    "Evaluates farmer welfare and bank risk jointly, recognizing that lower agricultural borrowing can reflect improved liquidity rather than an adverse contraction in credit supply.",
    "Uses OBBBA as a forward-looking case study of a large support expansion and labels the results as predictions rather than evidence of realized post-OBBBA effects.",
]:
    add_bullet(text)

doc.add_heading("4. Four Transmission Channels and Counterfactuals", level=1)
channels = [
    ("B1: Deposit-funding channel. ", "Farm-support payments may remain temporarily in local deposits. Banks then choose whether to replace wholesale funding, expand agricultural lending, purchase securities, or retain liquidity. This channel activates the estimated deposit response while holding borrower demand, default risk, and government credit unchanged."),
    ("B2: Borrower-liquidity channel. ", "Because farmers receive the payments directly, greater liquidity may reduce demand for short-term operating loans, permit debt repayment, finance consumption or investment internally, and improve future borrowing capacity. This channel activates the historical agricultural-loan-demand response while holding the other channels unchanged."),
    ("B3: Default-risk channel. ", "Government support may improve repayment capacity and reduce delinquencies, charge-offs, expected loan losses, and banks' marginal lending costs. The historical regression does not causally identify this effect, so the central calibration assumes a 10 percent reduction in charge-offs and reports 0 and 25 percent sensitivity cases."),
    ("B4: Government-credit channel. ", "Higher Marketing Assistance Loan rates improve farmers' government-supported financing option. This may improve liquidity and bargaining power while substituting for commercial operating loans. The magnitude is calibrated because the project does not yet estimate substitution between bank loans and Marketing Assistance Loans."),
]
for prefix, body in channels:
    add_bullet(prefix + body, bold_prefix=prefix)

doc.add_heading("Counterfactual design", level=2)
counterfactuals = [
    ("A_no_policy. ", "Baseline with all OBBBA policy shocks set to zero under current market power."),
    ("B_obbba_current_market_power. ", "Primary counterfactual with all four channels operating simultaneously under the estimated deposit and agricultural-loan market power."),
    ("B1-B4 channel decompositions. ", "Activate each transmission channel separately to identify which mechanism drives the full-policy prediction."),
    ("B5_obbba_no_default_effect. ", "Full policy with no reduction in default risk."),
    ("B6_obbba_chargeoffs_down_25pct. ", "Full policy with a stronger 25 percent charge-off reduction."),
]
for prefix, body in counterfactuals:
    add_bullet(prefix + body, bold_prefix=prefix)

comparison = doc.add_paragraph()
comparison.add_run("Primary comparison. ").bold = True
comparison.add_run(
    "B_obbba_current_market_power minus A_no_policy. B1-B4 explain the mechanisms; "
    "B5-B6 test uncertainty about default risk and are not separate government programs."
)
shade_paragraph(comparison, fill="F4F6F9")

doc.add_heading("5. Datasets Downloaded and Constructed", level=1)
datasets = [
    ("FSA annual payment files, 2014-2025. ", "104 official recipient/payment workbooks (approximately 3.5 GB). Recipient records were aggregated to county, year, and program; names and addresses remain outside Git. Programs were classified as ARC/PLC, MFP, CFAP, other FSA, and total FSA payments."),
    ("FSA ARC/PLC program data. ", "Official 2025 county-crop enrolled base acres, 2025 county PLC yields, projected 2026 effective prices and PLC payment rates, 2026 market-year-average price information, and 2026 ARC-CO county benchmark revenues."),
    ("BEA county farm-income data. ", "CAINC45 county government payments and farm/personal income. The banking overlap is 1994-2022 because BEA discontinued the detailed CAINC45 table after 2022; it was not artificially extended."),
    ("FDIC Summary of Deposits, 1994-2025. ", "Branch-level deposits were aggregated to county markets and used to construct bank service-area weights from prior-year branch deposit shares."),
    ("FDIC Call Reports, 1994-2025. ", "Bank deposits, agricultural production loans, total agricultural loans, total loans, equity, capital ratios, net charge-offs, interest income and expense, reserves, wholesale/nondeposit funding, repos, dividends, taxes, employment, premises expense, and other balance-sheet variables."),
    ("USDA ERS agricultural finance data. ", "Annual agricultural debt by lender, including commercial banks, the Farm Credit System, the Farm Service Agency, and other lenders; these series define the outside-credit environment."),
    ("USDA ERS farm income and wealth data. ", "National net cash farm income was converted to real terms and used to construct farm-income innovations and strong, normal, and downturn states."),
    ("Federal Reserve/FRED data. ", "Daily effective federal funds rate observations were aggregated to annual funding-rate states for the dynamic bank model."),
    ("USDA ERS Marketing Assistance Loan rates. ", "Pre-OBBBA and OBBBA statutory commodity loan rates were recorded for barley, corn, soybeans, wheat, rice, pulses, cotton, peanuts, and other covered crops. Most rates increase by about 10 percent, with heterogeneous increases across commodities."),
    ("Constructed analytical panels. ", "County payment-retention panel, bank service-area policy-exposure panel, deposit/total-loan/agricultural-production-loan BLP panels, dynamic bank-model panel, OBBBA county payment simulations, bank deposit shocks, default sensitivities, and Marketing Assistance Loan exposure indices."),
]
for prefix, body in datasets:
    add_bullet(prefix + body, bold_prefix=prefix)

doc.add_heading("6. Methods", level=1)
methods = [
    ("Payment timing and geography. ", "FSA disbursements from July through December are assigned to the following June SOD observation; January through June payments are assigned to the same June. County payments are not assigned directly to individual banks. Bank exposure is a service-area measure based on prior-year SOD county deposit weights."),
    ("Payment-retention regression. ", "County deposit growth is regressed on total FSA payments divided by lagged county deposits, with county and year fixed effects and standard errors clustered by county. Separate ARC/PLC, MFP, and CFAP regressions are retained as diagnostics."),
    ("OBBBA payment simulation. ", "Baseline and OBBBA ARC/PLC payments are calculated from county-crop base acres, PLC yields, projected effective prices, reference prices, ARC benchmarks, guarantee rates, payment caps, and a proportional scenario for the additional 30 million base acres. ARC outcomes are evaluated over revenue scenarios from 70 to 100 percent of benchmark revenue."),
    ("Deposit counterfactual. ", "The simulated increase in total FSA payments is multiplied by the statistically significant total-FSA deposit-retention coefficient. County deposit growth is mapped to banks using branch-deposit geography."),
    ("BLP market-power estimation. ", "Separate differentiated-product demand systems are estimated for deposits, total loans, and agricultural production loans. Bank and year fixed effects are included; salary expense/assets and premises expense/assets serve as the primary Wang-style supply-cost instruments. Demand derivatives are combined with ownership matrices to recover markups or markdowns."),
    ("Agricultural-lending reduced form. ", "Bank agricultural-loan balances and rates are related to total-FSA payment intensity in the bank's service area using bank and year fixed effects. This estimates the historical net association but does not separately identify borrower demand, bank supply, default risk, or program targeting."),
    ("Balance-sheet channel regressions. ", "Net charge-offs, wholesale funding, reserves, capital, agricultural-loan markups, deposit markdowns, and deposit rates are related to total-FSA service-area intensity. Effects that are not statistically identified are not treated as causal estimates."),
    ("Default-risk sensitivity. ", "Because historical payments are targeted to weak agricultural conditions, the charge-off regression cannot identify the causal insurance effect. The Bellman exercise therefore reports no default effect, a 10 percent charge-off reduction, and a 25 percent charge-off reduction."),
    ("Government-credit substitution. ", "OBBBA Marketing Assistance Loan rate changes are mapped to county crop exposure using enrolled base acres as a proxy. Because MAL eligibility follows current production rather than base acres, this is an exposure calibration, not a forecast of government loan volume."),
    ("Corrected Bellman model. ", "Banks choose loan rates, deposit rates, dividends, securities, reserves, and wholesale funding subject to deposit and loan demand, defaults, balance-sheet feasibility, retained earnings, reserve requirements, and capital constraints. The invalid earlier competition shortcut that multiplied demand slopes by 100 was removed."),
    ("Structural counterfactuals. ", "The model solves no policy; full OBBBA under current market power; funding-only, borrower-liquidity-only, default-only, and government-credit-only channels; and alternative default assumptions."),
    ("Market-power treatment. ", "Every reported policy counterfactual retains the deposit and agricultural-loan market power estimated in the data. Competitive-pricing calibrations are excluded. The current implementation uses a two-point bank-state grid and reports convergence diagnostics."),
]
for prefix, body in methods:
    add_bullet(prefix + body, bold_prefix=prefix)

doc.add_heading("7. Current Predictions", level=1)
results = [
    ("BLP market power. ", "The median deposit markdown is approximately 16.1 basis points; the median total-loan markup is 100.9 basis points; and the median agricultural-production-loan markup is 156.0 basis points."),
    ("Total-FSA deposit response. ", "The total-FSA coefficient is 0.07616 (standard error 0.02337; p = 0.0011). Interpreted mechanically, one additional dollar of FSA disbursement is associated with about 7.6 cents of additional county deposit stock. The separate ARC/PLC, MFP, and CFAP coefficients are not statistically significant and are not used in the main simulation."),
    ("Central OBBBA payment simulation. ", "At actual county revenue equal to 85 percent of ARC benchmark revenue, incremental OBBBA support is approximately $9.32 billion, a 61.2 percent increase relative to the 2025 total-FSA baseline of $15.23 billion."),
    ("Predicted deposits. ", "Applying the total-FSA coefficient predicts approximately $710 million in additional county deposits. The bank service-area mapping produces approximately $729 million because county growth rates are applied to banks' full deposit bases."),
    ("Historical agricultural lending response. ", "Total-FSA service-area intensity is associated with lower agricultural-production loan balances (coefficient -0.7618; p = 0.0098). The agricultural-loan-rate coefficient is positive but statistically insignificant (p = 0.232), so the project does not conclude that OBBBA raises loan rates."),
    ("Reduced-form OBBBA lending prediction. ", "The historical lending relationship predicts agricultural lending declines of approximately 1.21 percent at agricultural banks and 0.23 percent at nonagricultural banks, about $597 million combined. This is consistent with payments reducing farmers' short-term operating-credit needs."),
    ("Default risk. ", "The historical charge-off coefficient is not statistically identified (p = 0.188). The model therefore treats lower default risk as a sensitivity rather than an estimated effect. The 10 percent and 25 percent cases reduce average charge-offs by approximately 8.2 and 20.4 basis points, respectively."),
    ("Funding and capital. ", "Total-FSA exposure is associated with lower wholesale funding (coefficient -0.0335; p = 0.065) and higher bank capital (coefficient 0.0390; p < 0.001). This supports a funding-substitution channel but does not imply that all new deposits become loans."),
    ("Combined Bellman calibration under current market power. ", "For agricultural banks, the combined policy calibration reduces new agricultural loans by 6.56 percent and lowers loan rates by 14.6 basis points. For nonagricultural banks, new agricultural loans decline by 9.28 percent and the loan-rate change is approximately zero."),
    ("Farmer policy implication. ", "OBBBA can improve farmer liquidity and repayment capacity while simultaneously reducing demand for short-term commercial operating loans. Whether farmers also receive cheaper or more available bank credit depends on default-risk pass-through, commercial-bank market power, capital constraints, and competition from FSA/Farm Credit and Marketing Assistance Loans."),
    ("Interpretation limitation. ", "The Bellman accounting identities hold to numerical precision and no solved policy violates the imposed capital constraint. However, competitor-rate fixed-point gaps remain about 0.002-0.012. The structural magnitudes are calibration results and require a finer grid, full equilibrium convergence, and valid instruments for default and outside-credit elasticities before publication."),
]
for prefix, body in results:
    add_bullet(prefix + body, bold_prefix=prefix)

doc.add_heading("8. Selected References", level=1)
references = [
    "Giri, A. K., Subedi, D., and Kassel, K. (2022). Analysis of the payments from the Coronavirus Food Assistance Program and the Market Facilitation Program to minority producers. Applied Economic Perspectives and Policy. https://doi.org/10.1002/aepp.13325",
    "Goodwin, B. K., Mishra, A. K., and Ortalo-Magne, F. N. (2003). What's wrong with our models of agricultural land values? American Journal of Agricultural Economics, 85(3), 744-752. https://doi.org/10.1111/1467-8276.00479",
    "Hennessy, D. A. (1998). The production effects of agricultural income support policies under uncertainty. American Journal of Agricultural Economics, 80(1), 46-57.",
    "Janzen, J. P., Malone, T., Schaefer, K. A., and Scheitrum, D. P. (2023). Political returns to ad hoc farm payments? Applied Economic Perspectives and Policy, 45(1), 555-578. https://doi.org/10.1002/aepp.13216",
    "Janzen, J. P., Swearingen, B., and Yu, J. (2023). Buying time: The effect of Market Facilitation Program payments on the supply of grain storage. Journal of the Agricultural and Applied Economics Association, 2(3), 370-385. https://doi.org/10.1002/jaa2.67",
    "Key, N. (2020). Off-farm income, credit constraints, and farm investment. Journal of Agricultural and Applied Economics, 52(4), 642-663.",
    "Kirwan, B. E. (2009). The incidence of U.S. agricultural subsidies on farmland rental rates. Journal of Political Economy, 117(1), 138-164. https://doi.org/10.1086/598688",
    "Roberts, M. J., Kirwan, B., and Hopkins, J. (2003). The incidence of government program payments on agricultural land rents: The challenges of identification. American Journal of Agricultural Economics, 85(3), 762-769. https://doi.org/10.1111/1467-8276.00481",
    "Sun, Z. (manuscript). Government-support program announcements and corn and soybean futures volatility: Market Facilitation Program and Coronavirus Food Assistance Program.",
    "Walraven, N., and Barry, P. J. (2003). Bank risk ratings and the pricing of agricultural loans. Federal Reserve Board Finance and Economics Discussion Series.",
    "Wang, Y., Whited, T. M., Wu, Y., and Xiao, K. (2020). Bank market power and monetary policy transmission: Evidence from a structural estimation. SSRN 3049665.",
]
for ref in references:
    p = doc.add_paragraph(ref)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(5)

doc.core_properties.title = "How Government Farm Support Transmits Through Banks"
doc.core_properties.subject = "Farmer welfare, bank risk, literature, methods, and OBBBA predictions"
doc.core_properties.author = "AG Banking Research Project"

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
